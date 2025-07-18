"""
SOP Document Processing Service
Handles extraction, processing, and analysis of Standard Operating Procedures documents
"""

import os
import logging
import mimetypes
from typing import Dict, Any, List, Optional, Tuple
import tempfile
import asyncio
from datetime import datetime
from pathlib import Path

# Document processing libraries
try:
    import pypdf
    import docx
    from pdfplumber import PDF
    import textract
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")

# AI processing
try:
    import openai
    # Mock Agent to avoid import issues
    class MockAgent:
        def __init__(self, model_name: str, system_prompt: str = "", deps_type=None):
            self.model_name = model_name
            self.system_prompt = system_prompt
            self.deps_type = deps_type
            self.tools = []
        
        def tool(self, func):
            self.tools.append(func)
            return func

    # Use mock class instead of pydantic_ai import
    Agent = MockAgent
except ImportError as e:
    logging.warning(f"AI processing libraries not available: {e}")

from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

class SOPProcessor:
    """Service for processing SOP documents and extracting training-relevant content."""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
            'text/csv': 'csv'
        }
        
        # Configure AI agent for content analysis
        self.content_analyzer = self._create_content_analyzer()
        
        # Processing limits
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.max_pages = 500
        
    def _create_content_analyzer(self):
        """Create AI agent for analyzing SOP content."""
        
        system_prompt = """
        You are an expert SOP (Standard Operating Procedures) content analyzer specializing in extracting training-relevant information for English language learning programs.

        Your task is to analyze SOP documents and identify:
        1. **Key Vocabulary**: Industry-specific terms, technical language, professional expressions
        2. **Communication Patterns**: Common phrases, formal language structures, interaction protocols
        3. **Training Scenarios**: Real-world situations that require English communication
        4. **Skill Requirements**: Language skills needed for different roles and procedures
        5. **Cultural Context**: Professional norms, etiquette, and communication styles

        Extract content that would be valuable for creating customized English training programs for employees who need to follow these procedures.

        Focus on:
        - Technical terminology and jargon
        - Communication touchpoints with colleagues, clients, or external parties
        - Written communication requirements (emails, reports, documentation)
        - Verbal communication needs (meetings, presentations, phone calls)
        - Safety and compliance language
        - Quality control and standards terminology

        Provide structured, actionable insights that can inform curriculum design and content creation.
        """
        
        try:
            return Agent(
                'openai:gpt-4o-mini',
                system_prompt=system_prompt
            )
        except Exception as e:
            logger.warning(f"Could not initialize AI content analyzer: {e}")
            return None
    
    async def process_sop_files(self, files: List[UploadFile], course_request_id: int) -> Dict[str, Any]:
        """Process a list of SOP files and extract training-relevant content."""
        
        results = {
            "course_request_id": course_request_id,
            "processed_files": [],
            "extraction_summary": {},
            "training_insights": {},
            "vocabulary_analysis": {},
            "communication_patterns": {},
            "processing_errors": [],
            "total_processing_time": 0,
            "processed_at": datetime.utcnow().isoformat()
        }
        
        start_time = datetime.utcnow()
        
        for file in files:
            try:
                logger.info(f"Processing SOP file: {file.filename}")
                
                # Validate file
                validation_result = await self._validate_file(file)
                if not validation_result["valid"]:
                    results["processing_errors"].append({
                        "file": file.filename,
                        "error": validation_result["error"]
                    })
                    continue
                
                # Extract content from file
                extraction_result = await self._extract_file_content(file)
                if not extraction_result["success"]:
                    results["processing_errors"].append({
                        "file": file.filename,
                        "error": extraction_result["error"]
                    })
                    continue
                
                # Analyze content for training insights
                analysis_result = await self._analyze_content_for_training(
                    extraction_result["content"], 
                    file.filename
                )
                
                # Store file processing result
                file_result = {
                    "filename": file.filename,
                    "file_type": extraction_result["file_type"],
                    "file_size": extraction_result["file_size"],
                    "page_count": extraction_result.get("page_count", 1),
                    "word_count": extraction_result.get("word_count", 0),
                    "content_preview": extraction_result["content"][:500] + "..." if len(extraction_result["content"]) > 500 else extraction_result["content"],
                    "training_analysis": analysis_result,
                    "processing_status": "success",
                    "processed_at": datetime.utcnow().isoformat()
                }
                
                results["processed_files"].append(file_result)
                
                # Aggregate insights
                if analysis_result.get("success", False):
                    self._aggregate_insights(results, analysis_result)
                
            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {e}")
                results["processing_errors"].append({
                    "file": file.filename,
                    "error": f"Processing failed: {str(e)}"
                })
        
        # Finalize results
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        results["total_processing_time"] = processing_time
        results["files_processed"] = len(results["processed_files"])
        results["files_failed"] = len(results["processing_errors"])
        results["success_rate"] = len(results["processed_files"]) / len(files) if files else 0
        
        logger.info(f"SOP processing completed: {results['files_processed']}/{len(files)} files successful")
        
        return results
    
    async def _validate_file(self, file: UploadFile) -> Dict[str, Any]:
        """Validate uploaded SOP file."""
        
        # Check file size
        if hasattr(file, 'size') and file.size > self.max_file_size:
            return {
                "valid": False,
                "error": f"File size ({file.size} bytes) exceeds maximum ({self.max_file_size} bytes)"
            }
        
        # Check file type
        mime_type = mimetypes.guess_type(file.filename)[0]
        if mime_type not in self.supported_formats:
            return {
                "valid": False,
                "error": f"Unsupported file type: {mime_type}. Supported types: {list(self.supported_formats.values())}"
            }
        
        # Basic filename validation
        if not file.filename or len(file.filename) > 255:
            return {
                "valid": False,
                "error": "Invalid filename"
            }
        
        return {"valid": True}
    
    async def _extract_file_content(self, file: UploadFile) -> Dict[str, Any]:
        """Extract text content from uploaded file."""
        
        try:
            # Read file content
            content = await file.read()
            file_size = len(content)
            
            # Reset file pointer for potential re-reading
            await file.seek(0)
            
            # Determine file type
            mime_type = mimetypes.guess_type(file.filename)[0]
            file_type = self.supported_formats.get(mime_type, 'unknown')
            
            # Extract text based on file type
            if file_type == 'pdf':
                text_content = await self._extract_pdf_content(content)
            elif file_type == 'docx':
                text_content = await self._extract_docx_content(content)
            elif file_type == 'txt':
                text_content = content.decode('utf-8', errors='ignore')
            else:
                # Fallback: try textract for other formats
                text_content = await self._extract_with_textract(content, file.filename)
            
            # Basic content analysis
            word_count = len(text_content.split()) if text_content else 0
            
            return {
                "success": True,
                "content": text_content,
                "file_type": file_type,
                "file_size": file_size,
                "word_count": word_count,
                "mime_type": mime_type
            }
            
        except Exception as e:
            logger.error(f"Content extraction failed for {file.filename}: {e}")
            return {
                "success": False,
                "error": f"Content extraction failed: {str(e)}"
            }
    
    async def _extract_pdf_content(self, content: bytes) -> str:
        """Extract text content from PDF file."""
        
        try:
            import io
            
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            # Check page count limit
            if len(pdf_reader.pages) > self.max_pages:
                logger.warning(f"PDF has {len(pdf_reader.pages)} pages, limiting to first {self.max_pages}")
            
            text_content = ""
            for i, page in enumerate(pdf_reader.pages):
                if i >= self.max_pages:
                    break
                text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            # Fallback to pdfplumber if available
            try:
                import pdfplumber
                import io
                
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text_content = ""
                    for i, page in enumerate(pdf.pages):
                        if i >= self.max_pages:
                            break
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                    return text_content.strip()
                    
            except Exception as e2:
                logger.error(f"pdfplumber extraction also failed: {e2}")
                raise Exception(f"Could not extract PDF content: {e}")
    
    async def _extract_docx_content(self, content: bytes) -> str:
        """Extract text content from DOCX file."""
        
        try:
            import io
            import docx
            
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise Exception(f"Could not extract DOCX content: {e}")
    
    async def _extract_with_textract(self, content: bytes, filename: str) -> str:
        """Extract content using textract library as fallback."""
        
        try:
            import textract
            import tempfile
            
            # Write content to temporary file
            with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            try:
                # Extract text using textract
                text_content = textract.process(temp_path).decode('utf-8', errors='ignore')
                return text_content.strip()
            finally:
                # Clean up temporary file
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"Textract extraction error: {e}")
            raise Exception(f"Could not extract content with textract: {e}")
    
    async def _analyze_content_for_training(self, content: str, filename: str) -> Dict[str, Any]:
        """Analyze extracted content for training insights using AI."""
        
        if not self.content_analyzer:
            # Fallback to basic analysis without AI
            return await self._basic_content_analysis(content, filename)
        
        try:
            # Prepare analysis prompt
            analysis_prompt = f"""
            Analyze this SOP document content for English language training purposes:

            Document: {filename}
            Content Preview: {content[:2000]}...

            Please provide:
            1. Key vocabulary and technical terms
            2. Communication scenarios and touchpoints
            3. Required language skills and proficiency levels
            4. Suggested training focus areas
            5. Industry-specific language patterns

            Format your response as structured JSON with clear categories.
            """
            
            # Run AI analysis
            result = await self.content_analyzer.run(analysis_prompt)
            
            # Parse and structure the AI response
            analysis_result = {
                "success": True,
                "analysis_method": "ai_powered",
                "filename": filename,
                "insights": result.data if hasattr(result, 'data') else str(result),
                "analyzed_at": datetime.utcnow().isoformat()
            }
            
            # Try to extract structured data from AI response
            try:
                import json
                if isinstance(result.data, str):
                    parsed_insights = json.loads(result.data)
                    analysis_result["structured_insights"] = parsed_insights
            except:
                # If parsing fails, keep original format
                pass
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"AI analysis failed for {filename}: {e}")
            # Fallback to basic analysis
            return await self._basic_content_analysis(content, filename)
    
    async def _basic_content_analysis(self, content: str, filename: str) -> Dict[str, Any]:
        """Basic content analysis without AI (fallback method)."""
        
        try:
            # Basic text analysis
            words = content.lower().split()
            word_count = len(words)
            unique_words = len(set(words))
            
            # Identify potential technical terms (words > 6 characters, containing specific patterns)
            technical_terms = set()
            for word in words:
                if len(word) > 6 and any(pattern in word for pattern in ['tion', 'ment', 'ness', 'ical', 'ology']):
                    technical_terms.add(word)
            
            # Look for communication keywords
            communication_keywords = [
                'email', 'meeting', 'report', 'presentation', 'call', 'discuss',
                'communicate', 'inform', 'notify', 'update', 'feedback', 'review'
            ]
            
            found_communication = [kw for kw in communication_keywords if kw in content.lower()]
            
            # Look for procedure keywords
            procedure_keywords = [
                'step', 'process', 'procedure', 'guideline', 'instruction', 'protocol',
                'standard', 'requirement', 'compliance', 'quality', 'safety'
            ]
            
            found_procedures = [kw for kw in procedure_keywords if kw in content.lower()]
            
            return {
                "success": True,
                "analysis_method": "basic_text_analysis",
                "filename": filename,
                "insights": {
                    "word_count": word_count,
                    "unique_words": unique_words,
                    "technical_terms": list(technical_terms)[:20],  # Limit to top 20
                    "communication_indicators": found_communication,
                    "procedure_indicators": found_procedures,
                    "complexity_score": min(unique_words / word_count * 100, 100) if word_count > 0 else 0
                },
                "analyzed_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Basic analysis failed for {filename}: {e}")
            return {
                "success": False,
                "error": f"Analysis failed: {str(e)}",
                "filename": filename
            }
    
    def _aggregate_insights(self, results: Dict[str, Any], analysis_result: Dict[str, Any]):
        """Aggregate insights from individual file analysis into overall results."""
        
        try:
            insights = analysis_result.get("insights", {})
            
            # Initialize aggregation structures if not exists
            if "vocabulary_themes" not in results["vocabulary_analysis"]:
                results["vocabulary_analysis"]["vocabulary_themes"] = set()
                results["vocabulary_analysis"]["technical_terms"] = set()
                results["vocabulary_analysis"]["complexity_scores"] = []
            
            if "communication_scenarios" not in results["communication_patterns"]:
                results["communication_patterns"]["communication_scenarios"] = set()
                results["communication_patterns"]["procedure_types"] = set()
            
            # Aggregate vocabulary data
            if isinstance(insights, dict):
                # Handle AI-generated insights
                if "technical_terms" in insights:
                    if isinstance(insights["technical_terms"], list):
                        results["vocabulary_analysis"]["technical_terms"].update(insights["technical_terms"])
                
                if "complexity_score" in insights:
                    results["vocabulary_analysis"]["complexity_scores"].append(insights["complexity_score"])
                
                # Handle communication data
                if "communication_indicators" in insights:
                    results["communication_patterns"]["communication_scenarios"].update(insights["communication_indicators"])
                
                if "procedure_indicators" in insights:
                    results["communication_patterns"]["procedure_types"].update(insights["procedure_indicators"])
            
            # Convert sets to lists for JSON serialization
            for key in results["vocabulary_analysis"]:
                if isinstance(results["vocabulary_analysis"][key], set):
                    results["vocabulary_analysis"][key] = list(results["vocabulary_analysis"][key])
            
            for key in results["communication_patterns"]:
                if isinstance(results["communication_patterns"][key], set):
                    results["communication_patterns"][key] = list(results["communication_patterns"][key])
            
        except Exception as e:
            logger.error(f"Error aggregating insights: {e}")
    
    async def get_processing_status(self, course_request_id: int) -> Dict[str, Any]:
        """Get the processing status for SOP files of a course request."""
        
        # This would typically query a database for processing status
        # For now, return a placeholder
        return {
            "course_request_id": course_request_id,
            "processing_status": "completed",
            "files_processed": 0,
            "processing_time": 0,
            "last_updated": datetime.utcnow().isoformat()
        }

# Global instance
sop_processor = SOPProcessor()