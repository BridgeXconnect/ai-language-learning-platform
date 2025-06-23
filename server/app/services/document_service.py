"""
Document processing service for handling SOP files and extracting content.
Supports PDF, DOCX, and TXT files.
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile
import hashlib

try:
    import PyPDF2
    from docx import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LangchainDocument
except ImportError as e:
    logging.warning(f"Document processing dependencies not installed: {e}")
    PyPDF2 = None
    Document = None
    RecursiveCharacterTextSplitter = None
    LangchainDocument = None

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of various document formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.text_splitter = None
        
        if RecursiveCharacterTextSplitter:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported."""
        return Path(filename).suffix.lower() in self.supported_formats
    
    async def process_file(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """Process a file and extract text content."""
        
        if not self.is_supported_format(file_path):
            raise ValueError(f"Unsupported file format. Supported: {self.supported_formats}")
        
        try:
            # If file_content is provided, write to temp file
            if file_content:
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file_path).suffix) as tmp_file:
                    tmp_file.write(file_content)
                    temp_path = tmp_file.name
                
                try:
                    result = await self._extract_text(temp_path)
                finally:
                    os.unlink(temp_path)  # Clean up temp file
            else:
                result = await self._extract_text(file_path)
            
            # Add metadata
            result['file_name'] = Path(file_path).name
            result['file_size'] = len(file_content) if file_content else os.path.getsize(file_path)
            result['content_hash'] = hashlib.md5(result['text'].encode()).hexdigest()
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise
    
    async def _extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from file based on format."""
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_ext == '.docx':
            return await self._extract_docx_text(file_path)
        elif file_ext == '.txt':
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
    
    async def _extract_pdf_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        
        if not PyPDF2:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        text = ""
        metadata = {"pages": 0, "format": "pdf"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                # Try to get PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', '')
                    })
        
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
        
        return {
            "text": text.strip(),
            "metadata": metadata,
            "word_count": len(text.split()),
            "char_count": len(text)
        }
    
    async def _extract_docx_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        
        if not Document:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        text = ""
        metadata = {"format": "docx", "paragraphs": 0}
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    metadata["paragraphs"] += 1
            
            # Extract tables
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text += " | ".join(row_text) + "\n"
            
            if table_text:
                text += "\n--- Tables ---\n" + table_text
                metadata["tables"] = len(doc.tables)
            
            # Get document properties
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            })
        
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
        
        return {
            "text": text.strip(),
            "metadata": metadata,
            "word_count": len(text.split()),
            "char_count": len(text)
        }
    
    async def _extract_txt_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file."""
        
        text = ""
        metadata = {"format": "txt", "encoding": "utf-8"}
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        metadata["encoding"] = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                raise ValueError("Could not decode text file with any supported encoding")
            
            # Count lines
            metadata["lines"] = len(text.splitlines())
        
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise
        
        return {
            "text": text.strip(),
            "metadata": metadata,
            "word_count": len(text.split()),
            "char_count": len(text)
        }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """Split text into chunks for processing."""
        
        if not self.text_splitter:
            # Fallback simple chunking if langchain not available
            chunks = []
            words = text.split()
            
            for i in range(0, len(words), chunk_size - overlap):
                chunk_words = words[i:i + chunk_size]
                chunk_text = " ".join(chunk_words)
                
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": len(chunks),
                    "word_count": len(chunk_words),
                    "char_count": len(chunk_text)
                })
            
            return chunks
        
        # Use langchain text splitter
        try:
            # Update splitter parameters
            self.text_splitter.chunk_size = chunk_size
            self.text_splitter.chunk_overlap = overlap
            
            text_chunks = self.text_splitter.split_text(text)
            
            chunks = []
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    "text": chunk,
                    "chunk_index": i,
                    "word_count": len(chunk.split()),
                    "char_count": len(chunk)
                })
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            # Fallback to simple chunking
            return self.chunk_text(text, chunk_size, overlap)
    
    async def extract_key_terms(self, text: str, max_terms: int = 50) -> List[Dict[str, Any]]:
        """Extract key terms and phrases from text."""
        
        import re
        from collections import Counter
        
        # Simple keyword extraction (can be enhanced with NLP libraries)
        
        # Remove common words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
            'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 
            'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
            'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 
            'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'its',
            'our', 'their'
        }
        
        # Extract words and phrases
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Filter out stop words
        filtered_words = [word for word in words if word not in stop_words]
        
        # Count frequencies
        word_freq = Counter(filtered_words)
        
        # Extract phrases (2-3 words)
        sentences = re.split(r'[.!?]', text)
        phrases = []
        
        for sentence in sentences:
            sentence_words = re.findall(r'\b[a-zA-Z]+\b', sentence.lower())
            
            # Extract 2-word phrases
            for i in range(len(sentence_words) - 1):
                phrase = f"{sentence_words[i]} {sentence_words[i+1]}"
                if not any(word in stop_words for word in phrase.split()):
                    phrases.append(phrase)
            
            # Extract 3-word phrases
            for i in range(len(sentence_words) - 2):
                phrase = f"{sentence_words[i]} {sentence_words[i+1]} {sentence_words[i+2]}"
                if not any(word in stop_words for word in phrase.split()):
                    phrases.append(phrase)
        
        phrase_freq = Counter(phrases)
        
        # Combine and rank terms
        key_terms = []
        
        # Add top words
        for word, freq in word_freq.most_common(max_terms // 2):
            key_terms.append({
                "term": word,
                "frequency": freq,
                "type": "word",
                "relevance_score": freq / len(filtered_words)
            })
        
        # Add top phrases
        for phrase, freq in phrase_freq.most_common(max_terms // 2):
            if freq > 1:  # Only include phrases that appear multiple times
                key_terms.append({
                    "term": phrase,
                    "frequency": freq,
                    "type": "phrase",
                    "relevance_score": freq / len(phrases)
                })
        
        # Sort by relevance score
        key_terms.sort(key=lambda x: x["relevance_score"], reverse=True)
        
        return key_terms[:max_terms]
    
    async def analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure and content patterns."""
        
        lines = text.splitlines()
        
        analysis = {
            "total_lines": len(lines),
            "non_empty_lines": len([line for line in lines if line.strip()]),
            "average_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0,
            "sections": [],
            "has_numbered_lists": False,
            "has_bullet_points": False,
            "has_headers": False
        }
        
        # Detect sections (lines that look like headers)
        import re
        
        header_patterns = [
            r'^\d+\.\s+[A-Z]',  # "1. SECTION"
            r'^[A-Z][A-Z\s]{5,}$',  # "ALL CAPS HEADERS"
            r'^\d+\.\d+\s+',  # "1.1 Subsection"
            r'^[A-Z][a-z]+:$'  # "Section:"
        ]
        
        sections = []
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if line_stripped:
                for pattern in header_patterns:
                    if re.match(pattern, line_stripped):
                        sections.append({
                            "line_number": i + 1,
                            "text": line_stripped,
                            "type": "header"
                        })
                        analysis["has_headers"] = True
                        break
        
        analysis["sections"] = sections
        
        # Detect lists
        text_lower = text.lower()
        if re.search(r'^\d+\.', text, re.MULTILINE):
            analysis["has_numbered_lists"] = True
        
        if re.search(r'^\s*[•\-*]', text, re.MULTILINE):
            analysis["has_bullet_points"] = True
        
        return analysis

# Global instance
document_processor = DocumentProcessor()